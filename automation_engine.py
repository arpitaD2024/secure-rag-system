from secure_rag import SecureRAGAgent
from docx import Document
import faiss
import re


class RAGAutomationEngine:

    def __init__(self):
        self.agent = SecureRAGAgent(mode="groq")

    # EXTRACT DOCX 
    def extract_docx(self, file_bytes):
        path = "temp.docx"
        with open(path, "wb") as f:
            f.write(file_bytes)

        doc = Document(path)

        lines = []

        # paragraphs
        for p in doc.paragraphs:
            if p.text.strip():
                lines.append(p.text.strip())

        # tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    [cell.text.strip() for cell in row.cells if cell.text.strip()]
                )
                if row_text:
                    lines.append(row_text)

        return "\n".join(lines)

 
    def smart_chunk(self, text, max_words=120):

        lines = [l.strip() for l in text.split("\n") if l.strip()]

        chunks = []
        buffer = []

        for line in lines:

            if len(line.split()) <= 6:
                buffer.append(line)
                continue

            buffer.append(line)

            # flush when large enough
            if len(" ".join(buffer).split()) >= max_words:
                chunks.append(" ".join(buffer))
                buffer = []

        if buffer:
            chunks.append(" ".join(buffer))

        return chunks

    # UPLOAD DOCUMENT
    def upload_document(self, session_id, file_bytes):

        text = self.extract_docx(file_bytes)

        chunks = self.smart_chunk(text)

        if not chunks:
            return {"status": "error", "message": "Empty document"}

        emb = self.agent.embedder.encode(chunks).astype("float32")

        index = faiss.IndexFlatL2(emb.shape[1])
        index.add(emb)

        # session-based storage (no overwrite)
        self.agent.session_indexes[session_id] = index
        self.agent.session_chunks[session_id] = chunks

        return {"status": "uploaded", "chunks": len(chunks)}

    # QUERY 
    def query(self, session_id, question):

        res = self.agent.run(question, session_id=session_id)

        return {
            "answer": res["answer"],
            "attack_detected": res["attack_detected"],
            "attack_type": res["attack_type"],
            "score": res["score"]
        }